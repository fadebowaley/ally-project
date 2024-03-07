import os
import requests
from bs4 import BeautifulSoup
import re
from docx import Document

# URL to scrape
url = "https://www.uscis.gov/working-in-the-united-states/permanent-workers/employment-based-immigration-first-preference-eb-1"

# Desktop path to save the files
desktop_path = r"C:\Users\PMD - FEMI\OneDrive\Desktop\USCIS"

# Function to clean filename
def clean_filename(filename):
    return re.sub(r'[<>:"/\\|?*]', '_', filename)

# Function to scrape content and save to Word document
def scrape_and_save(url, desktop_path):
    # Sending a GET request to the URL
    response = requests.get(url)
    
    # Checking if the request was successful
    if response.status_code == 200:
        # Parsing the HTML content
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Finding the specified table with class 'dataTable'
        table = soup.find('table', class_='dataTable')
        
        # Checking if the table is found
        if table:
            # Extracting the caption
            caption = table.find('caption').get_text().strip()
            
            # Initialize a new Word document
            doc = Document()
            
            # Add caption to the document
            doc.add_heading(caption, level=1)
            
            # Extracting and adding content within the table to the document
            rows = table.find_all('tr')
            for row in rows:
                cells = row.find_all(['td', 'th'])
                row_content = [cell.get_text().strip() for cell in cells]
                # Add tabulated content as a row in the table
                table_row = doc.add_table(rows=1, cols=len(row_content))
                table_row.style = 'Table Grid'
                for i, cell_content in enumerate(row_content):
                    table_row.cell(0, i).text = cell_content
            
            # Extracting title for filename and cleaning it
            title = clean_filename(caption)
            filename = title + ".docx"
            
            # Creating the file path
            file_path = os.path.join(desktop_path, filename)
            
            # Save the document
            doc.save(file_path)
                
            print(f"Content saved to: {file_path}")
        else:
            print("Table not found.")
    else:
        print("Failed to retrieve content from the URL.")

# Call the function
scrape_and_save(url, desktop_path)
