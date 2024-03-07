import os
import requests
from bs4 import BeautifulSoup
import re
from docx import Document
import warnings

# Suppress warnings
warnings.filterwarnings("ignore")

# URL to scrape
url = "https://www.uscis.gov/working-in-the-united-states/permanent-workers/employment-based-immigration-first-preference-eb-1"
# url = "https://www.uscis.gov/working-in-the-united-states/permanent-workers/employment-based-immigration-second-preference-eb-2"
# url = "https://www.uscis.gov/working-in-the-united-states/permanent-workers/employment-based-immigration-third-preference-eb-3"
#url = "https://www.uscis.gov/working-in-the-united-states/permanent-workers/employment-based-immigration-fourth-preference-eb-4"

# Desktop path to save the files
desktop_path = r"C:\Users\PMD - FEMI\OneDrive\Desktop\USCIS"

# Function to clean filename
def clean_filename(filename):
    return re.sub(r'[<>:"/\\|?*]', '_', filename)

# Function to scrape content and save to Word document
def scrape_and_save(url, desktop_path):
    # Sending a GET request to the URL
    response = requests.get(url)
    
    # Checking if the request was successful
    if response.status_code == 200:
        # Parsing the HTML content
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Finding paragraphs, <h2>, and <ul> within div with id 'page-content'
        page_content_div = soup.find('div', id='page-content')
        paragraphs = page_content_div.find_all(['p', 'h2', 'ul'])
        
        # Extracting the first paragraph before the table
        first_paragraph = paragraphs[0].get_text().strip()
        
        # Finding the specified table with class 'dataTable'
        table = soup.find('table', class_='dataTable')
        
        # Initialize a new Word document
        doc = Document()
        
        # Add the first paragraph to the document
        doc.add_paragraph(first_paragraph)
        
        # Checking if the table is found
        if table:
            # Extracting and adding content within the table to the document
            rows = table.find_all('tr')
            for row in rows:
                cells = row.find_all(['td', 'th'])
                row_content = [cell.get_text().strip() for cell in cells]
                # Add tabulated content as a row in the table
                table_row = doc.add_table(rows=1, cols=len(row_content))
                table_row.style = 'Table Grid'
                for i, cell_content in enumerate(row_content):
                    table_row.cell(0, i).text = cell_content
        else:
            print("Table not found.")
        
        # Add paragraphs, <h2>, and <ul> to the document
        for element in paragraphs[1:]:
            if element.name == 'p':
                doc.add_paragraph(element.get_text().strip())
            elif element.name == 'h2':
                doc.add_heading(element.get_text().strip(), level=2)
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.get_text().strip(), style='ListBullet')
        
        # Extracting title for filename and cleaning it
        title = clean_filename(url.split("/")[-1])
        filename = title + ".docx"
        
        # Creating the file path
        file_path = os.path.join(desktop_path, filename)
        
        # Save the document
        doc.save(file_path)
            
        print(f"Content saved to: {file_path}")
    else:
        print("Failed to retrieve content from the URL.")

# Call the function
scrape_and_save(url, desktop_path)
